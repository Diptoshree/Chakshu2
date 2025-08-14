
######################################zzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzzz

import os
import zipfile
import streamlit as st
from PIL import Image
from io import BytesIO
import google.generativeai as genai
from docx import Document
from docx.shared import Inches
import json
import re
api_key = st.secrets["secret_section"]["api_key"]
# Configure the client using the API key
genai.configure(api_key=api_key)
# ✅ Set your Gemini API key
#genai.configure(api_key="")  # Replace with your key

# 🔵 Original function for analyzing a newspaper image
def analyze_newspaper_image(image):
    prompt = """
You are an expert in analyzing newspaper clippings. From the newspaper image, extract the following structured information:
1. News Brand Name (e.g., 'Dainik Bhaskar', 'The Times of India')
2. Heading
3. Subheading (if any)
4. Callout Boxes (if any)(visually distinct boxed or highlighted texts. avoid adding any other words or explanation. Please keep the text as it is in original news.)
5. Date (if mentioned)
6. Location (usually before the main news)
7. News Bureau (e.g., 'New Delhi Bureau')
8. Body Text (remaining text that forms the main article body. Please avoid photo captions)
9. A concise summary of the article (within 120 words). Start the summary by repeating the original headline exactly, from next paragraph a brief summary in your own words. The language of summary should be same as the language of the original text.
10. Sentiment of the news from the point of view of welbeing the Cityzen/government/administration (The sentiment of the news in one word 'Positive' or 'Negative')
Return the result in this JSON format:
{
  "news_brand": "",
  "heading": "",
  "subheading": "",
  "callout_boxes": [],
  "date": "",
  "location": "",
  "news_bureau": "",
  "body_text": "",
  "summary": "",
  "sentiment": ""
}
"""
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content([prompt, image])
    return response.text

# 🔵 Format the extracted JSON into Word Document
def format_json_to_doc(doc, extracted_json_text):
    try:
        json_text_match = re.search(r"\{[\s\S]*\}", extracted_json_text)
        structured_data = json.loads(json_text_match.group())
    except:
        doc.add_paragraph("❌ Failed to parse structured JSON from Gemini response.")
        return

    doc.add_heading("📰 Extracted Newspaper Article", level=1)

    def add_section(label, content, style="Normal"):
        if content:
            doc.add_heading(label, level=2)
            if isinstance(content, list):
                for item in content:
                    doc.add_paragraph(f"• {item}", style=style)
            else:
                doc.add_paragraph(content.strip(), style=style)

    add_section("Heading", structured_data.get("heading"))
    add_section("Sentiment", structured_data.get("sentiment"))
    add_section("Summary", structured_data.get("summary"))

    heading = structured_data.get("heading")
    subheading = structured_data.get("subheading")
    body_text = structured_data.get("body_text")
    if any([heading, subheading, body_text]):
        doc.add_heading("Body Text", level=2)
        if heading:
            doc.add_paragraph(heading.strip())
        if subheading:
            doc.add_paragraph(subheading.strip())
        if body_text:
            doc.add_paragraph(body_text.strip())

    add_section("Callout Boxes", structured_data.get("callout_boxes", []))
    add_section("News Brand", structured_data.get("news_brand"))

# 🔵 Save single image and extracted text into Word
def save_to_word_with_image(image_file, extracted_text, doc=None):
    if doc is None:
        doc = Document()
    try:
        img = Image.open(image_file)
        buffer = BytesIO()
        img.save(buffer, format="JPEG")
        buffer.seek(0)
        doc.add_picture(buffer, width=Inches(5.5))
        doc.add_paragraph("⬆️ Original Newspaper Image", style="Caption")
    except Exception as e:
        doc.add_paragraph(f"❌ Failed to insert image: {e}")

    format_json_to_doc(doc, extracted_text)
    return doc

# 🔵 Streamlit UI
st.set_page_config(page_title="Newspaper Analyzer", layout="centered")
st.title("📰 Newspaper Image Analyzer")

# 👉 Mode Selection
mode = st.radio("Choose Mode", ("Single Image Search", "Bulk Image Search"))

if mode == "Single Image Search":
    uploaded_file = st.file_uploader("Upload Newspaper Image", type=["jpg", "jpeg", "png"])

    if uploaded_file:
        image = Image.open(uploaded_file)
        st.image(image, caption="🖼️ Uploaded Newspaper", use_container_width=True)

        if st.button("🔍 Analyze Image"):
            with st.spinner("Analyzing..."):
                extracted_text = analyze_newspaper_image(image)
                try:
                    json_text_match = re.search(r"\{[\s\S]*\}", extracted_text)
                    structured = json.loads(json_text_match.group())
                except:
                    structured = None

            if structured:
                st.success("✅ Extraction complete")
                st.subheader("📌 Extracted Details")
                st.markdown(f"**Headline:** {structured.get('heading', 'N/A')}")
                st.markdown(f"**Sentiment:** {structured.get('sentiment', 'N/A')}")
                st.markdown(f"**Summary:** {structured.get('summary', 'N/A')}")
                
                st.subheader("Body Text")
                if structured.get("heading"):
                    st.markdown(f"Headline: {structured['heading']}")
                if structured.get("subheading"):
                    st.markdown(f"Subheading: {structured['subheading']}")
                if structured.get("body_text"):
                    st.markdown(structured["body_text"])
                st.markdown(f"News Brand: {structured.get('news_brand', 'N/A')}")
                st.markdown(f"Callout Boxes: {structured.get('callout_boxes', '[]')}")

                word_file = BytesIO()
                save_to_word_with_image(uploaded_file, extracted_text).save(word_file)
                word_file.seek(0)
                st.download_button("📥 Download Word File", word_file, file_name="Extracted_News.docx")
            else:
                st.error("❌ Could not parse structured response.")

elif mode == "Bulk Image Search":
    uploaded_zip = st.file_uploader("Upload ZIP of Newspaper Images", type=["zip"])

    if uploaded_zip:
        if st.button("🔍 Analyze Bulk Images"):
            with st.spinner("Analyzing images..."):
                temp_folder = "temp_images"
                os.makedirs(temp_folder, exist_ok=True)

                # Extract ZIP
                with zipfile.ZipFile(uploaded_zip, "r") as zip_ref:
                    zip_ref.extractall(temp_folder)

                image_files = [os.path.join(temp_folder, f) for f in os.listdir(temp_folder)
                               if f.lower().endswith((".jpg", ".jpeg", ".png"))]

                if not image_files:
                    st.error("❌ No valid image files found in ZIP.")
                else:
                    doc = Document()
                    for img_path in image_files:
                        try:
                            img = Image.open(img_path)
                            extracted_text = analyze_newspaper_image(img)
                            save_to_word_with_image(img_path, extracted_text, doc)
                        except Exception as e:
                            doc.add_paragraph(f"⚠️ Failed to process {img_path}: {e}")

                    # Finalize Word file
                    output = BytesIO()
                    doc.save(output)
                    output.seek(0)

                    st.success(f"✅ Processed {len(image_files)} images.")
                    st.download_button("📥 Download Combined Word File", output, file_name="Bulk_Extracted_News.docx")

                # Cleanup temp folder
                try:
                    for f in os.listdir(temp_folder):
                        os.remove(os.path.join(temp_folder, f))
                    os.rmdir(temp_folder)
                except Exception as cleanup_error:
                    print(f"Cleanup error: {cleanup_error}")


