

import streamlit as st
import google.generativeai as genai
from PIL import Image, UnidentifiedImageError
import logging
import io
import zipfile
from docx import Document
from docx.shared import Inches
import base64

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')


api_key = st.secrets["secret_section"]["api_key"]
# Configure the client using the API key
genai.configure(api_key=api_key)

def extract_text_from_image(image):
    logging.info("Processing uploaded image")
    model = genai.GenerativeModel("gemini-2.5-pro-exp-03-25")
    prompt = """
    You are extracting text from a scanned newspaper cutting. Separate the text into two categories: HEADLINE and BODY TEXT.
    
    - **HEADLINE**: The main news headline, usually in the largest bold font, spanning between one to three lines. Ignore subheaders, captions, dates, and newspaper organization names.
    - **BODY TEXT**: All other content, including subheaders, captions, and article text.
    
    **Output format (strictly follow this structure without any labels, explanations, or additional formatting):**
    
    HEADLINE:
    [Extracted headline text]
    
    BODY TEXT:
    [Extracted body text]
    """
    
    response = model.generate_content([image, prompt])
    extracted_text = response.text.strip()
    print("Raw Extracted Text:", extracted_text)
    
    # Extract headline and body text based on explicit markers
    if "HEADLINE:" in extracted_text and "BODY TEXT:" in extracted_text:
        parts = extracted_text.split("BODY TEXT:")
        headline = parts[0].replace("HEADLINE:", "").strip()
        body_text = parts[1].strip() if len(parts) > 1 else "No additional text found"
    else:
        headline = "Headline not detected"
        body_text = "No additional text found"
    
    logging.info(f"Extracted Headline: {headline}")
    return headline, body_text

def generate_summary(article_text):
    logging.info("Generating summary...")
    summarization_prompt = f"""
    Summarize the following news article in less than {min(len(article_text.split()) // 2, 100)} words.
    Maintain the original language. The summary should include at least 3 key takeaways.
    Avoid adding a separate headline and structure it into two paragraphs.
    
    ARTICLE:
    {article_text}
    """
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(summarization_prompt)
    summary = response.text.strip()
    logging.info("Summary generated successfully.")
    return summary

def analyze_sentiment(article_text):
    logging.info("Performing sentiment analysis...")
    sentiment_prompt = f"""
    Analyze the sentiment of the following news article and return only one word: Positive or Negative.
    
    ARTICLE:
    {article_text}
    """
    model = genai.GenerativeModel("gemini-2.0-flash")
    response = model.generate_content(sentiment_prompt)
    sentiment = response.text.strip()
    logging.info(f"Sentiment detected: {sentiment}")
    return sentiment

def save_to_word(image, headline, sentiment, summary, body_text, doc):
    logging.info("Saving data to Word document...")
    img_buffer = io.BytesIO()
    image.save(img_buffer, format='PNG')
    img_buffer.seek(0)
    doc.add_picture(img_buffer, width=Inches(4.5))
    doc.add_heading("Headline:", level=1)
    doc.add_paragraph(headline)
    doc.add_heading("Sentiment:", level=1)
    doc.add_paragraph(sentiment)
    doc.add_heading("Summary:", level=1)
    doc.add_paragraph(summary)
    doc.add_heading("Extracted Text:", level=1)
    doc.add_paragraph(body_text)

def process_zip_file(zip_file):
    logging.info("Processing ZIP file...")
    doc = Document()
    with zipfile.ZipFile(zip_file, 'r') as archive:
        for file_name in archive.namelist():
            if file_name.lower().endswith((".png", ".jpg", ".jpeg")):
                with archive.open(file_name) as file:
                    try:
                        image = Image.open(file).convert("RGB")
                        headline, body_text = extract_text_from_image(image)
                        summary = generate_summary(body_text)
                        sentiment = analyze_sentiment(body_text)
                        save_to_word(image, headline, sentiment, summary, body_text, doc)
                    except UnidentifiedImageError:
                        logging.error(f"Skipping non-image file: {file_name}")
    output_filename = "Extracted_Text.docx"
    doc.save(output_filename)
    logging.info(f"Word document saved as: {output_filename}")
    return output_filename

def set_background(image_file):
    with open(image_file, "rb") as f:
        base64_image = base64.b64encode(f.read()).decode()
    background_css = f"""
    <style>
    .stApp {{
        background: url("data:image/jpg;base64,{base64_image}") no-repeat center center fixed;
        background-size: cover;
    }}
    </style>
    """
    st.markdown(background_css, unsafe_allow_html=True)

# Set the background image
set_background("newsback.jpg")

# Create a left vertical pane (Sidebar)
with st.sidebar:
    st.image("logo2.jpg", width=300)  # Adjust the width as needed
    st.write("""
The 'Chakshu News Summarizer' automates the  process  of news insights extraction from images by providing:
1. Headline Extraction
2. Sentiment Analysis
3. News Summary
4. Raw Text Extraction from Images
5. Option to Download the Output """)  # Additional info

st.title("Chakshu News Summarizer")
st.write("Upload an image or a ZIP file containing images to extract information.")
option = st.radio("Choose an upload option:", ["Single Image", "ZIP File"])

if option == "Single Image":
    uploaded_file = st.file_uploader("Choose an image", type=["jpg", "png", "jpeg"])
    if uploaded_file is not None:
        image = Image.open(uploaded_file).convert("RGB")
        st.image(image, caption="Uploaded Image", use_container_width=True)
        if st.button("Process Image"):
            headline, body_text = extract_text_from_image(image)
            summary = generate_summary(body_text)
            sentiment = analyze_sentiment(body_text)
            st.subheader("Extracted Headline")
            st.write(headline)
            st.subheader("Sentiment Analysis")
            st.write(sentiment)
            st.subheader("Summary")
            st.write(summary)
            doc = Document()
            save_to_word(image, headline, sentiment, summary, body_text, doc)
            doc.save("Extracted_Text.docx")
            with open("Extracted_Text.docx", "rb") as f:
                st.download_button("Download Word Document", f, file_name="Extracted_Text.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

elif option == "ZIP File":
    uploaded_zip = st.file_uploader("Choose a ZIP file", type=["zip"])
    if uploaded_zip is not None:
        if st.button("Process ZIP File"):
            word_file = process_zip_file(uploaded_zip)
            with open(word_file, "rb") as f:
                st.download_button("Download Word Document", f, file_name=word_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
